import json
import os
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import google.generativeai as genai

# 환경 변수에서 구글 API 키 로드 (없을 경우를 대비해 스크립트 실행 전 세팅 필요)
# 주의: 이 파일은 파이썬 백엔드에서 실행되므로 .env 파일을 읽거나 직접 키를 주입해야 합니다.
from dotenv import load_dotenv
load_dotenv('data-intrix-web/.env.local')

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
else:
    print("경고: GEMINI_API_KEY가 설정되지 않았습니다. 테스트 모드로 실행합니다.")

def get_gemini_response(prompt):
    if not GEMINI_API_KEY:
        return "[테스트 모드] API 키가 없어 임시 텍스트를 출력합니다. 실제 환경에서는 여기에 AI가 작성한 분석 글이 들어갑니다."
    try:
        model = genai.GenerativeModel('gemini-2.5-flash')
        response = model.generate_content(prompt)
        return response.text.replace('**', '') # 워드에 넣을 때 마크다운 별표 제거
    except Exception as e:
        print(f"Gemini API 오류: {e}")
        return "AI 분석 결과를 불러오는 중 오류가 발생했습니다."

def create_word_report(json_path, output_docx_path):
    print("1. 데이터 및 AI 프롬프트 분석 시작...")
    
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    stats_str = json.dumps(data, ensure_ascii=False, indent=2)
    
    # AI에게 분석 요청 (각 파트별 분리 요청으로 퀄리티 향상)
    prompt_exec = f"""당신은 최고급 비즈니스 데이터 애널리스트입니다. 
다음 데이터 {stats_str} 를 바탕으로 리포트 표지에 들어갈 [Executive Summary(핵심 요약본)]을 4~5줄로 작성해주세요.
매출에 가장 큰 영향을 주는 변수와 전체적인 비즈니스 건전성을 전문가 톤으로 요약해야 합니다."""
    exec_summary = get_gemini_response(prompt_exec)

    prompt_monthly = f"다음 데이터 {stats_str} 를 바탕으로, '월별 총매출 추이'에 대한 기술적이고 객관적인 캡션 설명을 2~3줄로 작성해주세요. 명령조(~해라, ~분석하라)를 절대 쓰지 말고, 평어체(~입니다, ~보여줍니다)로 현재의 상승/하락 흐름만 설명하세요."
    monthly_comment = get_gemini_response(prompt_monthly)

    prompt_category = f"다음 데이터 {stats_str} 를 바탕으로, '상위 메뉴별 총매출' 현황에 대한 기술적이고 객관적인 캡션 설명을 2~3줄로 작성해주세요. 명령조(~해라, ~분석하라)를 절대 쓰지 말고, 평어체(~입니다, ~차지합니다)로 캐시카우 메뉴의 중요도만 설명하세요."
    category_comment = get_gemini_response(prompt_category)
    
    prompt_regression = f"""당신은 자영업 매출 데이터 분석가입니다.
아래 통계 분석 결과를 바탕으로 반드시 지정 형식대로만 출력하세요. 서술형 문단 절대 금지.

통계 데이터: {stats_str}

=== 반드시 이 형식 그대로 출력 ===
[강점] 소제목(10자이내)
- 항목명: 설명 1~2문장, 수치 인용 필수.
- 항목명: 설명 1~2문장, 수치 인용 필수.
- 항목명: 설명 1~2문장, 수치 인용 필수.

[약점] 소제목(10자이내)
- 항목명: 설명 1~2문장, 수치 인용 필수.
- 항목명: 설명 1~2문장, 수치 인용 필수.
- 항목명: 설명 1~2문장, 수치 인용 필수.

[개선 방향] 소제목(10자이내)
- 항목명: 실행 방안 1~2문장, 구체적으로.
- 항목명: 실행 방안 1~2문장, 구체적으로.
- 항목명: 실행 방안 1~2문장, 구체적으로.
=== 출력 끝 ===

규칙: [강점][약점][개선 방향] 3섹션만 출력. 섹션 외 다른 텍스트 없음. 각 섹션 불릿 3~4개. 항목명과 설명은 콜론(:)으로 구분."""
    regression_analysis = get_gemini_response(prompt_regression)

    print("2. 워드 문서(.docx) 조립 시작...")
    
    # 웹 UI에서 보여주기 위해 JSON 파일에 AI 분석 결과 추가 저장
    data['ai_insights'] = regression_analysis
    data['ai_monthly_caption'] = monthly_comment
    data['ai_category_caption'] = category_comment
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)
        
    document = Document()
    
    # 기본 폰트 설정 (맑은 고딕)
    style = document.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # ---------------------------------------------------------
    # 표지 구성 (First Page)
    # ---------------------------------------------------------
    # 제목
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("스파 매장 매출 분석 및 전략 리포트\n")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    
    # 날짜
    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f"작성일: {datetime.datetime.now().strftime('%Y년 %m월 %d일')}\n\n")
    date_run.font.size = Pt(12)
    
    # Executive Summary
    document.add_heading('Executive Summary (요약본)', level=1)
    document.add_paragraph(exec_summary)
    
    # 분석 방법 안내
    document.add_heading('분석 방법 안내', level=1)
    document.add_paragraph("본 리포트는 POS 매출 데이터를 기반으로 탐색적 데이터 분석(EDA) 및 변수 간 단순/다중회귀분석(LASSO 적용)을 수행하여 도출된 통계적 유의성을 바탕으로 AI 컨설팅 인텔리전스가 작성하였습니다.\n")
    
    # 소속
    affil_p = document.add_paragraph()
    affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil_p.add_run("\n\n\n\n\n[분석: 중간계 인트릭스 연구소]")
    affil_run.font.size = Pt(14)
    affil_run.font.bold = True
    
    document.add_page_break()

    # ---------------------------------------------------------
    # 본문 구성
    # ---------------------------------------------------------
    # 1. 월별 총매출 추이
    document.add_heading('매출 트렌드 시각화', level=1)
    if os.path.exists('output_charts/monthly_trend.png'):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture('output_charts/monthly_trend.png', width=Inches(6.0))
    
    p_comment = document.add_paragraph(monthly_comment)
    p_comment.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph("\n")

    # 2. 메뉴별 매출 현황
    document.add_heading('메뉴별 매출 데이터 분석 결과', level=1)
    if os.path.exists('output_charts/top_categories.png'):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture('output_charts/top_categories.png', width=Inches(6.0))
        
    p_comment2 = document.add_paragraph(category_comment)
    p_comment2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_page_break()

    # 3. 회귀분석 결과 표 및 심층 분석
    document.add_heading('회귀분석 결과 및 심화 분석', level=1)
    
    # 표 삽입 (단순회귀)
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '독립변수(메뉴)'
    hdr_cells[1].text = 'R-Squared (설명력)'
    hdr_cells[2].text = 'Coefficient (계수)'
    hdr_cells[3].text = 'P-value (유의확률)'
    
    for item in data.get('regression_simple', []):
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['메뉴변수'])
        row_cells[1].text = str(item['R_squared'])
        row_cells[2].text = str(item['계수(Coefficient)'])
        row_cells[3].text = str(item['P_value'])
        
    # AI 해석 삽입 (장점, 발견사항, 액션플랜 등)
    document.add_paragraph("\n")
    document.add_paragraph(regression_analysis)

    document.save(output_docx_path)
    print(f"3. 증권사 리포트 양식 워드 파일이 성공적으로 생성되었습니다: {output_docx_path}")

if __name__ == '__main__':
    import sys
    json_path = sys.argv[1] if len(sys.argv) > 1 else 'analysis_results.json'
    out_path = sys.argv[2] if len(sys.argv) > 2 else 'DataIntrix_Consulting_Report.docx'
    
    if os.path.exists(json_path):
        create_word_report(json_path, out_path)
    else:
        print(f"{json_path} 파일이 없습니다. 먼저 analyzer.py를 실행하세요.")
