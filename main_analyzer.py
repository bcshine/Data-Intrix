import pandas as pd
import numpy as np
import statsmodels.api as sm
from statsmodels.stats.outliers_influence import variance_inflation_factor
import json
import sys
import os
import path
from datetime import datetime
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# JSON 인코딩 에러 방지
class NumpyEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.integer): return int(obj)
        if isinstance(obj, np.floating): return float(obj)
        if isinstance(obj, np.ndarray): return obj.tolist()
        return super(NumpyEncoder, self).default(obj)

def run_analysis(input_csv):
    print(f"Starting analysis for {input_csv}...")
    
    # 1. 데이터 로드 및 전처리
    try:
        df = pd.read_csv(input_csv)
    except Exception as e:
        return {"error": f"CSV 로드 실패: {str(e)}"}

    # Wide/Long 자동 감지 및 변환
    if 'Category' in df.columns and 'Total_Amt' in df.columns:
        df = df.pivot_table(index='Period_Start', columns='Category', values='Total_Amt', aggfunc='sum').fillna(0).reset_index()
    
    if 'Period_Start' not in df.columns:
        return {"error": "데이터에 'Period_Start' 컬럼이 없습니다."}

    df = df.sort_values('Period_Start')
    numeric_df = df.select_dtypes(include=[np.number])
    target_col = 'Total_Sales' if 'Total_Sales' in numeric_df.columns else None
    
    if not target_col:
        df['Total_Sales'] = numeric_df.sum(axis=1)
        target_col = 'Total_Sales'
        numeric_df = df.select_dtypes(include=[np.number])

    # 2. 통계 분석 (상관관계, 회귀분석, 안정성)
    results = {}
    
    # 기초 통계 및 추세
    results['trend_analysis'] = {
        "추세_기울기": float(np.polyfit(range(len(df)), df[target_col], 1)[0]) if len(df) > 1 else 0
    }
    
    # 변동계수 (안정성)
    cv_list = []
    for col in numeric_df.columns:
        if col == target_col: continue
        mean = numeric_df[col].mean()
        std = numeric_df[col].std()
        cv = (std / mean * 100) if mean > 0 else 0
        cv_list.append({"메뉴": col, "평균": mean, "표준편차": std, "변동계수_CV_perc": round(cv, 1)})
    results['cv_stats'] = sorted(cv_list, key=lambda x: x['변동계수_CV_perc'])

    # 상관관계
    results['correlation_matrix'] = numeric_df.corr().to_dict()

    # 단순 회귀분석
    reg_results = []
    for col in numeric_df.columns:
        if col == target_col: continue
        try:
            X = sm.add_constant(numeric_df[col])
            model = sm.OLS(numeric_df[target_col], X).fit()
            reg_results.append({
                "메뉴변수": col,
                "회귀계수": model.params[1],
                "R_squared": model.rsquared,
                "P_value": model.pvalues[1]
            })
        except: continue
    results['regression_simple'] = sorted(reg_results, key=lambda x: x['R_squared'], reverse=True)

    # 3. AI 인사이트 생성 (Gemini)
    api_key = os.getenv("GOOGLE_API_KEY")
    if api_key:
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.0-flash-thinking-exp-1219')
            
            stats_summary = f"""
            - 분석대상: {input_csv}
            - 총 누적 매출: {df[target_col].sum():,.0f}원
            - 성장 추세: {'상승' if results['trend_analysis']['추세_기울기'] > 0 else '하락'}
            - 핵심 동인(R2 상위): {', '.join([r['메뉴변수'] for r in results['regression_simple'][:3]])}
            - 불안정 메뉴(CV 하위): {', '.join([r['메뉴'] for r in results['cv_stats'][-2:]])}
            """
            
            prompt = f"""
            스파/매장 매출 분석 전문가로서 다음 통계 데이터를 기반으로 경영 전략 리포트를 작성하라.
            
            [데이터 요약]
            {stats_summary}
            
            [작성 규칙]
            1. [강점], [약점], [개선 방향] 섹션으로 나누어 작성.
            2. 각 섹션 내에서는 '항목명: 설명' 구조의 불릿 포인트로 작성.
            3. 일반적인 이야기가 아닌, 제공된 '핵심 동인'과 '안정성' 데이터를 구체적으로 언급할 것.
            4. 전문적이고 통찰력 있는 증권사 리포트 톤앤매너 유지.
            """
            
            response = model.generate_content(prompt)
            results['ai_insights'] = response.text
            
            # 워드 리포트 생성
            create_word_report(df, results, "DataIntrix_Consulting_Report.docx")
            
        except Exception as e:
            results['ai_insights'] = f"AI 분석 중 오류 발생: {str(e)}"
    
    # 4. 결과 저장
    with open('analysis_results.json', 'w', encoding='utf-8') as f:
        json.dump(results, f, cls=NumpyEncoder, ensure_ascii=False, indent=2)
    
    print("Analysis complete. Results saved to analysis_results.json")
    return results

def create_word_report(df, stats, output_path):
    doc = Document()
    
    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)

    # 표지
    title = doc.add_heading('데이터 기반 매출 분석 및 전략 리포트', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"작성일: {datetime.now().strftime('%Y-%m-%d')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # AI 인사이트 섹션
    doc.add_heading('1. AI 전략적 제언', level=1)
    if 'ai_insights' in stats:
        doc.add_paragraph(stats['ai_insights'])
    
    # 통계 섹션
    doc.add_heading('2. 핵심 통계 요약', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '메뉴'
    hdr_cells[1].text = '평균 매출'
    hdr_cells[2].text = 'R-squared'
    hdr_cells[3].text = 'P-value'

    for r in stats['regression_simple'][:10]:
        row_cells = table.add_row().cells
        row_cells[0].text = r['메뉴변수']
        row_cells[1].text = f"{r.get('회귀계수', 0):,.0f}"
        row_cells[2].text = f"{r['R_squared']:.3f}"
        row_cells[3].text = f"{r['P_value']:.4f}"

    doc.save(output_path)
    print(f"Word report saved to {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python main_analyzer.py <csv_file>")
    else:
        run_analysis(sys.argv[1])
